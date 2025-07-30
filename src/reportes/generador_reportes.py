"""
Generador de Reportes para Análisis de Series Temporales

Clase principal para generar reportes automáticos en formatos PDF y Word
con gráficos, estadísticas y análisis embebidos.
"""

import os
import sys
from pathlib import Path
from datetime import datetime
import pandas as pd
import jinja2
from typing import Dict, List, Optional, Any

# Importar utilidades del módulo
from .utils_reportes import (
    exportar_grafico_plotly, 
    generar_estadisticas, 
    crear_grafico_resumen,
    formatear_numero
)

class GeneradorReportes:
    """
    Clase principal para generar reportes automáticos
    """
    
    def __init__(self, datos: pd.DataFrame, metadatos: pd.DataFrame):
        """
        Inicializa el generador de reportes
        
        Args:
            datos: DataFrame con los datos de las series temporales
            metadatos: DataFrame con metadatos de las series
        """
        self.datos = datos
        self.metadatos = metadatos
        
        # Solo generar estadísticas si hay datos válidos
        if not datos.empty and not metadatos.empty:
            self.estadisticas = generar_estadisticas(datos, metadatos)
        else:
            self.estadisticas = {
                'fecha_generacion': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                'total_series': 0,
                'total_datos': 0,
                'rango_fechas': {'inicio': 'N/A', 'fin': 'N/A'},
                'tipos_unicos': 0,
                'categorias_unicas': 0,
                'estadisticas_numericas': {
                    'valor_min': 0, 'valor_max': 0, 'valor_promedio': 0,
                    'valor_mediana': 0, 'desviacion_estandar': 0
                },
                'estadisticas_por_tipo': {},
                'estadisticas_por_categoria': {}
            }
        
        # Configurar Jinja2
        self.template_dir = Path(__file__).parent / 'templates'
        self.jinja_env = jinja2.Environment(
            loader=jinja2.FileSystemLoader(str(self.template_dir)),
            autoescape=True
        )
        
        # Registrar filtros personalizados
        self.jinja_env.filters['formatear_numero'] = formatear_numero
        self.jinja_env.globals['pd'] = pd
        
        # Crear estructura de directorios para mejor trazabilidad
        self.base_output_dir = Path('reportes_generados')
        self.base_output_dir.mkdir(exist_ok=True)
        
        # Crear subdirectorios por fecha
        fecha_actual = datetime.now().strftime('%Y-%m-%d')
        self.output_dir = self.base_output_dir / fecha_actual
        self.output_dir.mkdir(exist_ok=True)
        
        # Crear subdirectorios por tipo de reporte
        self.pdf_dir = self.output_dir / 'pdf'
        self.word_dir = self.output_dir / 'word'
        self.html_dir = self.output_dir / 'html'
        
        for dir_path in [self.pdf_dir, self.word_dir, self.html_dir]:
            dir_path.mkdir(exist_ok=True)
    
    def generar_pdf(self, 
                   graficos_especificos: Optional[Dict[str, Any]] = None,
                   nombre_archivo: Optional[str] = None) -> str:
        """
        Genera un reporte en formato PDF
        
        Args:
            graficos_especificos: Diccionario con gráficos adicionales
            nombre_archivo: Nombre del archivo de salida
        
        Returns:
            str: Ruta del archivo PDF generado
        """
        try:
            # Verificar si WeasyPrint está disponible
            try:
                from weasyprint import HTML, CSS
                from weasyprint.text.fonts import FontConfiguration
            except ImportError as e:
                raise ImportError(
                    f"WeasyPrint no está instalado. Instálalo con: pip install weasyprint. Error: {e}"
                )
            
            # Generar nombre de archivo si no se proporciona
            if nombre_archivo is None:
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                nombre_archivo = f"reporte_series_temporales_{timestamp}.pdf"
            
            ruta_pdf = self.pdf_dir / nombre_archivo
            
            # Crear gráfico de resumen
            grafico_resumen = crear_grafico_resumen(self.datos, self.metadatos)
            grafico_resumen_b64 = exportar_grafico_plotly(grafico_resumen)
            
            # Preparar datos para el template
            template_data = {
                'estadisticas': self.estadisticas,
                'metadatos': self.metadatos,
                'grafico_resumen': grafico_resumen_b64,
                'graficos_especificos': graficos_especificos or {}
            }
            
            # Renderizar template HTML
            template = self.jinja_env.get_template('template_pdf.html')
            html_content = template.render(**template_data)
            
            # Configurar fuentes
            font_config = FontConfiguration()
            
            # Generar PDF
            HTML(string=html_content).write_pdf(
                ruta_pdf,
                font_config=font_config
            )
            
            print(f"✅ Reporte PDF generado: {ruta_pdf}")
            return str(ruta_pdf)
            
        except Exception as e:
            print(f"❌ Error al generar PDF: {e}")
            raise
    
    def generar_word(self, 
                    graficos_especificos: Optional[Dict[str, Any]] = None,
                    nombre_archivo: Optional[str] = None) -> str:
        """
        Genera un reporte en formato Word (.docx)
        
        Args:
            graficos_especificos: Diccionario con gráficos adicionales
            nombre_archivo: Nombre del archivo de salida
        
        Returns:
            str: Ruta del archivo Word generado
        """
        try:
            # Verificar si python-docx está disponible
            try:
                from docx import Document
                from docx.shared import Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.oxml.shared import OxmlElement, qn
            except ImportError as e:
                raise ImportError(
                    f"python-docx no está instalado. Instálalo con: pip install python-docx. Error: {e}"
                )
            
            # Generar nombre de archivo si no se proporciona
            if nombre_archivo is None:
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                nombre_archivo = f"reporte_series_temporales_{timestamp}.docx"
            
            ruta_word = self.word_dir / nombre_archivo
            
            # Crear documento
            doc = Document()
            
            # Título principal
            title = doc.add_heading('📊 Reporte de Análisis de Series Temporales', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Información de generación
            doc.add_paragraph(f"Generado el: {self.estadisticas['fecha_generacion']}")
            doc.add_paragraph(f"Total de series analizadas: {self.estadisticas['total_series']}")
            doc.add_paragraph(f"Total de puntos de datos: {self.estadisticas['total_datos']}")
            
            # Resumen ejecutivo
            doc.add_heading('📋 Resumen Ejecutivo', level=1)
            resumen = doc.add_paragraph()
            resumen.add_run(
                f"Este reporte presenta un análisis completo de {self.estadisticas['total_series']} "
                f"series temporales, cubriendo un período desde {self.estadisticas['rango_fechas']['inicio']} "
                f"hasta {self.estadisticas['rango_fechas']['fin']}. "
                f"Se analizaron {self.estadisticas['total_datos']} puntos de datos distribuidos en "
                f"{self.estadisticas['tipos_unicos']} tipos y {self.estadisticas['categorias_unicas']} categorías diferentes."
            )
            
            # Estadísticas generales
            doc.add_heading('📈 Estadísticas Generales', level=1)
            
            # Crear tabla de estadísticas
            stats_table = doc.add_table(rows=1, cols=2)
            stats_table.style = 'Table Grid'
            
            # Encabezados
            hdr_cells = stats_table.rows[0].cells
            hdr_cells[0].text = 'Métrica'
            hdr_cells[1].text = 'Valor'
            
            # Agregar estadísticas
            stats_data = [
                ('Total de Series', str(self.estadisticas['total_series'])),
                ('Total de Datos', str(self.estadisticas['total_datos'])),
                ('Tipos Únicos', str(self.estadisticas['tipos_unicos'])),
                ('Categorías Únicas', str(self.estadisticas['categorias_unicas'])),
                ('Valor Promedio', formatear_numero(self.estadisticas['estadisticas_numericas']['valor_promedio'])),
                ('Desviación Estándar', formatear_numero(self.estadisticas['estadisticas_numericas']['desviacion_estandar'])),
                ('Valor Mínimo', formatear_numero(self.estadisticas['estadisticas_numericas']['valor_min'])),
                ('Valor Máximo', formatear_numero(self.estadisticas['estadisticas_numericas']['valor_max']))
            ]
            
            for metric, value in stats_data:
                row_cells = stats_table.add_row().cells
                row_cells[0].text = metric
                row_cells[1].text = value
            
            # Estadísticas por tipo
            doc.add_heading('📊 Estadísticas por Tipo', level=1)
            
            tipo_table = doc.add_table(rows=1, cols=6)
            tipo_table.style = 'Table Grid'
            
            # Encabezados
            hdr_cells = tipo_table.rows[0].cells
            hdr_cells[0].text = 'Tipo'
            hdr_cells[1].text = 'Cantidad'
            hdr_cells[2].text = 'Promedio'
            hdr_cells[3].text = 'Desv. Estándar'
            hdr_cells[4].text = 'Mínimo'
            hdr_cells[5].text = 'Máximo'
            
            # Agregar datos por tipo
            for tipo, stats in self.estadisticas['estadisticas_por_tipo'].items():
                row_cells = tipo_table.add_row().cells
                row_cells[0].text = str(tipo)
                row_cells[1].text = str(stats['count'])
                row_cells[2].text = formatear_numero(stats['mean'])
                row_cells[3].text = formatear_numero(stats['std'])
                row_cells[4].text = formatear_numero(stats['min'])
                row_cells[5].text = formatear_numero(stats['max'])
            
            # Metadatos de las series
            doc.add_heading('📋 Metadatos de las Series', level=1)
            
            meta_table = doc.add_table(rows=1, cols=6)
            meta_table.style = 'Table Grid'
            
            # Encabezados
            hdr_cells = meta_table.rows[0].cells
            hdr_cells[0].text = 'ID Serie'
            hdr_cells[1].text = 'Tipo'
            hdr_cells[2].text = 'Categoría'
            hdr_cells[3].text = 'Unidad'
            hdr_cells[4].text = 'Fecha Inicio'
            hdr_cells[5].text = 'Fecha Fin'
            
            # Agregar metadatos
            for _, serie in self.metadatos.iterrows():
                row_cells = meta_table.add_row().cells
                row_cells[0].text = str(serie['id_serie'])
                row_cells[1].text = str(serie['tipo'])
                row_cells[2].text = str(serie['categoria'])
                row_cells[3].text = str(serie['unidad'])
                row_cells[4].text = serie['fecha_inicio'].strftime('%Y-%m-%d') if pd.notna(serie['fecha_inicio']) else 'N/A'
                row_cells[5].text = serie['fecha_fin'].strftime('%Y-%m-%d') if pd.notna(serie['fecha_fin']) else 'N/A'
            
            # Guardar documento
            doc.save(ruta_word)
            
            print(f"✅ Reporte Word generado: {ruta_word}")
            return str(ruta_word)
            
        except Exception as e:
            print(f"❌ Error al generar Word: {e}")
            raise
    
    def generar_html(self, 
                    graficos_especificos: Optional[Dict[str, Any]] = None,
                    nombre_archivo: Optional[str] = None) -> str:
        """
        Genera un reporte en formato HTML
        
        Args:
            graficos_especificos: Diccionario con gráficos adicionales
            nombre_archivo: Nombre del archivo de salida
        
        Returns:
            str: Ruta del archivo HTML generado
        """
        try:
            # Generar nombre de archivo si no se proporciona
            if nombre_archivo is None:
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                nombre_archivo = f"reporte_series_temporales_{timestamp}.html"
            
            ruta_html = self.html_dir / nombre_archivo
            
            # Crear gráfico de resumen
            grafico_resumen = crear_grafico_resumen(self.datos, self.metadatos)
            grafico_resumen_b64 = exportar_grafico_plotly(grafico_resumen)
            
            # Preparar datos para el template
            template_data = {
                'estadisticas': self.estadisticas,
                'metadatos': self.metadatos,
                'grafico_resumen': grafico_resumen_b64,
                'graficos_especificos': graficos_especificos or {}
            }
            
            # Renderizar template HTML
            template = self.jinja_env.get_template('template_pdf.html')
            html_content = template.render(**template_data)
            
            # Guardar archivo HTML
            with open(ruta_html, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            print(f"✅ Reporte HTML generado: {ruta_html}")
            return str(ruta_html)
            
        except Exception as e:
            print(f"❌ Error al generar HTML: {e}")
            raise
    
    def generar_todos_formatos(self, 
                              graficos_especificos: Optional[Dict[str, Any]] = None,
                              prefijo: Optional[str] = None) -> Dict[str, str]:
        """
        Genera reportes en todos los formatos disponibles
        
        Args:
            graficos_especificos: Diccionario con gráficos adicionales
            prefijo: Prefijo para los nombres de archivo
        
        Returns:
            Dict[str, str]: Diccionario con las rutas de los archivos generados
        """
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        prefijo = prefijo or "reporte_series_temporales"
        
        archivos_generados = {}
        
        # Generar HTML (siempre disponible)
        try:
            nombre_html = f"{prefijo}_{timestamp}.html"
            archivos_generados['html'] = self.generar_html(
                graficos_especificos, nombre_html
            )
        except Exception as e:
            print(f"⚠️ No se pudo generar HTML: {e}")
        
        # Generar Word
        try:
            nombre_word = f"{prefijo}_{timestamp}.docx"
            archivos_generados['word'] = self.generar_word(
                graficos_especificos, nombre_word
            )
        except Exception as e:
            print(f"⚠️ No se pudo generar Word: {e}")
        
        # Generar PDF
        try:
            nombre_pdf = f"{prefijo}_{timestamp}.pdf"
            archivos_generados['pdf'] = self.generar_pdf(
                graficos_especificos, nombre_pdf
            )
        except Exception as e:
            print(f"⚠️ No se pudo generar PDF: {e}")
        
        return archivos_generados
    
    def listar_reportes_generados(self) -> Dict[str, List[str]]:
        """
        Lista todos los reportes generados organizados por fecha
        
        Returns:
            Dict[str, List[str]]: Diccionario con fechas y rutas de archivos
        """
        reportes = {}
        
        if not self.base_output_dir.exists():
            return reportes
        
        # Recorrer todas las carpetas de fechas
        for fecha_dir in sorted(self.base_output_dir.iterdir(), reverse=True):
            if fecha_dir.is_dir():
                fecha_str = fecha_dir.name
                archivos_fecha = []
                
                # Buscar archivos en subdirectorios
                for subdir in ['pdf', 'word', 'html']:
                    subdir_path = fecha_dir / subdir
                    if subdir_path.exists():
                        for archivo in subdir_path.glob('*'):
                            if archivo.is_file():
                                archivos_fecha.append(str(archivo))
                
                if archivos_fecha:
                    reportes[fecha_str] = sorted(archivos_fecha)
        
        return reportes
    
    def obtener_estadisticas_reportes(self) -> Dict[str, Any]:
        """
        Obtiene estadísticas sobre los reportes generados
        
        Returns:
            Dict[str, Any]: Estadísticas de reportes
        """
        reportes = self.listar_reportes_generados()
        
        total_archivos = sum(len(archivos) for archivos in reportes.values())
        total_fechas = len(reportes)
        
        # Contar por tipo
        tipos = {'pdf': 0, 'word': 0, 'html': 0}
        for fecha, archivos in reportes.items():
            for archivo in archivos:
                if archivo.endswith('.pdf'):
                    tipos['pdf'] += 1
                elif archivo.endswith('.docx'):
                    tipos['word'] += 1
                elif archivo.endswith('.html'):
                    tipos['html'] += 1
        
        return {
            'total_archivos': total_archivos,
            'total_fechas': total_fechas,
            'tipos': tipos,
            'fechas_disponibles': list(reportes.keys())
        } 